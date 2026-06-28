"""
Resume parsing service using PyMuPDF (PDF) and python-docx (DOCX).
Stores embeddings in ChromaDB.
"""
import asyncio
import hashlib
import io
import json
import logging
import re
from pathlib import Path
from typing import Optional

import aiofiles
from langchain_core.messages import HumanMessage, SystemMessage
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.db.models import Resume
from app.services.ai.llm_provider import invoke_with_fallback

logger = logging.getLogger(__name__)


def _compute_file_hash(contents: bytes) -> str:
    """Compute SHA-256 hash of file contents for duplicate detection."""
    return hashlib.sha256(contents).hexdigest()


def _validate_pdf_magic_bytes(contents: bytes) -> bool:
    """Check that file actually starts with %PDF- header."""
    return contents[:5] == b"%PDF-"


def _validate_docx_magic_bytes(contents: bytes) -> bool:
    """DOCX is a ZIP archive — check for PK header."""
    return contents[:2] == b"PK"


async def parse_resume_file(file_path: str, mime_type: str) -> str:
    """Extract raw text from PDF or DOCX file. Runs parsing in a thread pool."""
    path = Path(file_path)

    if mime_type == "application/pdf" or path.suffix.lower() == ".pdf":
        return await asyncio.to_thread(_parse_pdf_sync, file_path)
    elif mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ) or path.suffix.lower() in (".docx", ".doc"):
        return await asyncio.to_thread(_parse_docx_sync, file_path)
    else:
        raise ValueError(f"Unsupported file type: {mime_type}")


def _parse_pdf_sync(file_path: str) -> str:
    """Synchronous PDF parsing using pypdf — called via asyncio.to_thread."""
    from pypdf import PdfReader

    try:
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            text_parts.append(page.extract_text() or "")
        return "\n".join(text_parts)
    except Exception as e:
        logger.error("PDF parsing error: %s", e)
        raise ValueError(f"Could not parse PDF: {e}")


def _parse_docx_sync(file_path: str) -> str:
    """Synchronous DOCX parsing — called via asyncio.to_thread."""
    from docx import Document

    try:
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error("DOCX parsing error: %s", e)
        raise ValueError(f"Could not parse DOCX: {e}")


async def extract_structured_data(raw_text: str) -> dict:
    """
    Use LLM to extract structured data from raw resume text.
    Returns a dict matching ParsedResumeData schema.
    """
    # Truncate to avoid token limits
    truncated_text = raw_text[:4000]

    prompt = f"""Extract structured information from this resume text. Return ONLY valid JSON.

RESUME TEXT:
{truncated_text}

Return this JSON structure (use empty lists/null if info not found):
{{
    "name": "Full Name or null",
    "email": "email@example.com or null",
    "phone": "phone number or null",
    "linkedin": "LinkedIn URL or null",
    "github": "GitHub URL or null",
    "skills": ["skill1", "skill2"],
    "technical_skills": ["Python", "React", "SQL"],
    "soft_skills": ["Leadership", "Communication"],
    "education": [
        {{
            "institution": "University Name",
            "degree": "B.Tech / B.S. / M.S.",
            "field": "Computer Science",
            "year": "2023",
            "gpa": "3.8 or null"
        }}
    ],
    "experience": [
        {{
            "company": "Company Name",
            "role": "Software Engineer",
            "duration": "2021-2023",
            "responsibilities": ["Did X", "Built Y"],
            "technologies": ["Python", "React"]
        }}
    ],
    "projects": [
        {{
            "name": "Project Name",
            "description": "Brief description",
            "technologies": ["tech1", "tech2"],
            "url": "URL or null"
        }}
    ],
    "certifications": [
        {{
            "name": "AWS Solutions Architect",
            "issuer": "Amazon",
            "year": "2022"
        }}
    ],
    "total_years_experience": 2.5,
    "primary_tech_stack": ["Python", "React", "PostgreSQL"]
}}"""

    response = await invoke_with_fallback([
        SystemMessage(content="You are a resume parsing expert. Return only valid JSON, nothing else."),
        HumanMessage(content=prompt),
    ])

    # Parse JSON response
    try:
        text = response.strip()
        # Remove markdown code blocks if present
        text = re.sub(r"```(?:json)?\n?", "", text).strip()
        return json.loads(text)
    except json.JSONDecodeError:
        # Try to extract JSON object
        match = re.search(r"\{.*\}", response, re.DOTALL)
        if match:
            try:
                return json.loads(match.group(0))
            except json.JSONDecodeError:
                pass

    logger.warning("Could not parse structured data from resume, returning basic extraction")
    return _basic_extraction(raw_text)


def _basic_extraction(text: str) -> dict:
    """Fallback basic regex extraction when LLM fails."""
    emails = re.findall(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b", text)
    phones = re.findall(r"[\+\(]?[1-9][0-9 .\-\(\)]{8,}[0-9]", text)
    skills_keywords = [
        "Python", "Java", "JavaScript", "TypeScript", "React", "Node.js",
        "SQL", "PostgreSQL", "MongoDB", "AWS", "Docker", "Kubernetes",
        "Machine Learning", "TensorFlow", "PyTorch", "FastAPI", "Django",
        "C++", "C#", "Go", "Rust", "Git", "Linux", "Redis", "GraphQL",
    ]
    found_skills = [s for s in skills_keywords if s.lower() in text.lower()]

    return {
        "name": None,
        "email": emails[0] if emails else None,
        "phone": phones[0] if phones else None,
        "linkedin": None,
        "github": None,
        "skills": found_skills,
        "technical_skills": found_skills,
        "soft_skills": [],
        "education": [],
        "experience": [],
        "projects": [],
        "certifications": [],
        "total_years_experience": None,
        "primary_tech_stack": found_skills[:5],
    }


async def generate_skill_summary(parsed_data: dict) -> str:
    """Generate a concise skill summary for the resume."""
    skills = parsed_data.get("technical_skills", []) + parsed_data.get("skills", [])
    skills = list(set(skills))[:20]
    experience = parsed_data.get("experience", [])
    projects = parsed_data.get("projects", [])

    prompt = f"""Write a 2-3 sentence professional skill summary based on this data.

Skills: {', '.join(skills)}
Experience count: {len(experience)} roles
Projects count: {len(projects)} projects
Years experience: {parsed_data.get('total_years_experience', 'unknown')}

Write a concise, professional summary starting with the candidate's level."""

    summary = await invoke_with_fallback([
        SystemMessage(content="You write concise professional summaries."),
        HumanMessage(content=prompt),
    ])
    return summary.strip()


async def analyze_missing_skills(parsed_data: dict, job_role: str) -> list[str]:
    """Identify skills typically required for the role that are missing from the resume."""
    existing_skills = set(s.lower() for s in parsed_data.get("skills", []) + parsed_data.get("technical_skills", []))

    prompt = f"""List the top 8 skills typically required for a {job_role} role that are NOT in this skill set.

Candidate's existing skills: {', '.join(list(existing_skills)[:20])}

Return ONLY a JSON array of skill strings:
["skill1", "skill2", "skill3"]"""

    response = await invoke_with_fallback([
        SystemMessage(content="You are a technical recruiter. Return only a JSON array."),
        HumanMessage(content=prompt),
    ])

    try:
        text = response.strip()
        match = re.search(r"\[.*\]", text, re.DOTALL)
        if match:
            return json.loads(match.group(0))
    except Exception:
        pass
    return []


async def store_resume_embedding(resume_id: str, text: str, metadata: dict) -> str:
    """Store resume text embedding in ChromaDB."""
    try:
        import chromadb
        client = chromadb.HttpClient(
            host=settings.CHROMA_HOST,
            port=settings.CHROMA_PORT,
        )
        collection = client.get_or_create_collection(
            name=settings.CHROMA_COLLECTION_RESUMES,
            metadata={"hnsw:space": "cosine"},
        )
        # Use resume_id as document ID
        collection.upsert(
            ids=[resume_id],
            documents=[text[:2000]],  # truncate for embedding
            metadatas=[{
                "user_id": metadata.get("user_id", ""),
                "filename": metadata.get("filename", ""),
                "job_role": metadata.get("job_role", ""),
            }],
        )
        return resume_id
    except Exception as e:
        logger.warning("ChromaDB embedding failed: %s — continuing without vector storage", e)
        return resume_id


async def get_resume_text(db: AsyncSession, resume_id: str) -> Optional[str]:
    """Retrieve parsed resume text from the database."""
    from sqlalchemy import select
    result = await db.execute(
        select(Resume).where(Resume.id == resume_id)
    )
    resume = result.scalar_one_or_none()
    return resume.raw_text if resume else None
